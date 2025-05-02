import os
import json
import logging
import re
from datetime import datetime
from flask import Flask, request, abort, jsonify
from dotenv import load_dotenv  
from linebot.v3 import WebhookHandler
from linebot.v3.messaging import Configuration, ApiClient, MessagingApi
from linebot.v3.webhooks import MessageEvent, TextMessageContent
from linebot.v3.messaging import (
    TextMessage, FlexMessage, FlexContainer, ReplyMessageRequest,
    QuickReply, QuickReplyItem, MessageAction
)
from linebot.v3.exceptions import InvalidSignatureError
from google.cloud.dialogflow_v2 import SessionsClient
from google.cloud.dialogflow_v2.types import TextInput, QueryInput
from google.protobuf.json_format import MessageToDict

# ไลบรารีเพิ่มเติมสำหรับอ่านไฟล์
import docx  # สำหรับอ่านไฟล์ .docx
from difflib import SequenceMatcher  # สำหรับเปรียบเทียบความคล้ายของข้อความ

# โหลด environment variables จากไฟล์ .env
load_dotenv()

# ตรวจสอบและกำหนดค่าตัวแปรสภาพแวดล้อมที่จำเป็น
GOOGLE_APPLICATION_CREDENTIALS = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
if not GOOGLE_APPLICATION_CREDENTIALS:
    raise ValueError("ไม่พบตัวแปรสภาพแวดล้อม GOOGLE_APPLICATION_CREDENTIALS")

DIALOGFLOW_PROJECT_ID = os.getenv("DIALOGFLOW_PROJECT_ID")
if not DIALOGFLOW_PROJECT_ID:
    raise ValueError("ไม่พบตัวแปรสภาพแวดล้อม DIALOGFLOW_PROJECT_ID")

LINE_CHANNEL_ACCESS_TOKEN = os.getenv("LINE_CHANNEL_ACCESS_TOKEN")
if not LINE_CHANNEL_ACCESS_TOKEN:
    raise ValueError("ไม่พบตัวแปรสภาพแวดล้อม LINE_CHANNEL_ACCESS_TOKEN")

LINE_CHANNEL_SECRET = os.getenv("LINE_CHANNEL_SECRET")
if not LINE_CHANNEL_SECRET:
    raise ValueError("ไม่พบตัวแปรสภาพแวดล้อม LINE_CHANNEL_SECRET")

# กำหนดค่า Config
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = GOOGLE_APPLICATION_CREDENTIALS
SESSION_ID = "line-bot-session"

# Flask App
app = Flask(__name__)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = app.logger

# Line Bot
configuration = Configuration(access_token=LINE_CHANNEL_ACCESS_TOKEN)
handler = WebhookHandler(LINE_CHANNEL_SECRET)
api_client = ApiClient(configuration)
line_bot_api = MessagingApi(api_client)

# โหลด Hugging Face model
huggingface_qa_pipeline = None
try:
    from transformers import pipeline
    # ใช้โมเดลภาษาไทย
    huggingface_qa_pipeline = pipeline("text-generation", model="flax-community/gpt2-base-thai")
    logger.info("โหลดโมเดล Hugging Face สำเร็จ")
except Exception as e:
    logger.error(f"ไม่สามารถโหลดโมเดล Hugging Face: {str(e)}")

# คำตอบที่ไม่ต้องการจาก Dialogflow (หากได้คำตอบเหล่านี้จะถือว่า Dialogflow ไม่สามารถตอบคำถามได้)
INVALID_DIALOGFLOW_RESPONSES = [
    "ขอโทษค่ะ พูดอีกครั้งได้ไหมคะ",
    "ขอโทษค่ะ ไม่เข้าใจ",
    "พูดใหม่อีกครั้งได้ไหมคะ",
    "ขอโทษค่ะ ฉันไม่เข้าใจค่ะ",
    "พูดอีกทีได้ไหมคะ" 
]

@app.route("/callback", methods=['POST'])
def callback():
    """
    เส้นทาง webhook สำหรับรับข้อความจาก LINE
    """
    signature = request.headers.get('X-Line-Signature', '')
    body = request.get_data(as_text=True)
    logger.info(f"ได้รับคำขอ: {body}")

    try:
        handler.handle(body, signature)
    except InvalidSignatureError:
        logger.error("ลายเซ็นไม่ถูกต้อง")
        abort(400)

    return 'OK'

@handler.add(MessageEvent, message=TextMessageContent)
def handle_message(event):
    """
    จัดการข้อความที่ได้รับจากผู้ใช้ LINE
    """
    user_id = event.source.user_id
    text_from_user = event.message.text
    logger.info(f"ข้อความจาก {user_id}: {text_from_user}")

    # ตรวจสอบว่าเป็นข้อความในกลุ่มหรือห้องสนทนาหรือไม่
    is_group = hasattr(event.source, 'type') and event.source.type in ['group', 'room']
    bot_name = "DPA Chatbot"
    should_respond = False
    actual_message = text_from_user

    # ถ้าเป็นข้อความในกลุ่ม ตรวจสอบว่าเรียกบอทหรือไม่
    if is_group and text_from_user.startswith(f'@{bot_name}'):
        actual_message = text_from_user.split(f'@{bot_name}', 1)[-1].strip()
        should_respond = True
        logger.info(f"ตรวจพบข้อความในกลุ่ม แยกข้อความได้: {actual_message}")
    elif not is_group:
        # ถ้าเป็นการแชทส่วนตัว ตอบทุกข้อความ
        should_respond = True

    if should_respond:
        if not actual_message:
            # กรณีที่ข้อความว่างเปล่า (เช่น เพียงแค่เรียกชื่อบอทในกลุ่ม)
            reply_text = f"สวัสดีค่ะ หนูชื่อ {bot_name} คุณต้องการสอบถามอะไรค่ะ?"
            send_text_message(event.reply_token, reply_text)
        else:
            try:
                # ขั้นตอนที่ 1: ส่งคำถามไปยัง Dialogflow
                logger.info("เริ่มขั้นตอนที่ 1: ส่งคำถามไปยัง Dialogflow")
                user_session_id = f"{SESSION_ID}-{user_id}"
                response = detect_intent_texts(DIALOGFLOW_PROJECT_ID, user_session_id, actual_message, 'th')
                
                # แปลง response เป็น dict สำหรับตรวจสอบ
                response_dict = MessageToDict(response._pb)
                logger.info(f"คำตอบจาก Dialogflow: {json.dumps(response_dict, indent=2, ensure_ascii=False)[:1000]}")
                
                # เตรียมข้อมูลที่จะส่งกลับ
                messages_to_reply = []
                quick_replies = []
                has_payload = False
                
                # ตรวจสอบและรวบรวมทุกข้อความจาก Dialogflow
                if 'queryResult' in response_dict:
                    # ตรวจสอบหา messages ที่มี quick replies และ payloads
                    if 'fulfillmentMessages' in response_dict['queryResult']:
                        for message in response_dict['queryResult']['fulfillmentMessages']:
                            # ตรวจสอบข้อความปกติ
                            if 'text' in message and 'text' in message['text'] and message['text']['text']:
                                for text in message['text']['text']:
                                    if text and text not in INVALID_DIALOGFLOW_RESPONSES:
                                        text_message = TextMessage(text=text)
                                        messages_to_reply.append(text_message)
                            
                            # ตรวจสอบ Quick Replies
                            if 'quickReplies' in message:
                                quick_reply_items = []
                                if 'quickReplies' in message['quickReplies'] and isinstance(message['quickReplies']['quickReplies'], list):
                                    for qr in message['quickReplies']['quickReplies']:
                                        quick_reply_items.append(
                                            QuickReplyItem(
                                                action=MessageAction(
                                                    label=qr[:20],  # LINE จำกัดความยาวของป้ายชื่อไม่เกิน 20 ตัวอักษร
                                                    text=qr
                                                )
                                            )
                                        )
                                if quick_reply_items:
                                    quick_replies = QuickReply(items=quick_reply_items)
                                    # ใส่ quick replies ในข้อความสุดท้าย (ถ้ามี)
                                    if messages_to_reply:
                                        messages_to_reply[-1].quick_reply = quick_replies
                                    
                            # ตรวจสอบ Custom Payload (เช่น Flex Message)
                            if 'payload' in message:
                                has_payload = True
                                process_payload(message['payload'], messages_to_reply)
                
                # ถ้าไม่มีข้อความใน fulfillmentMessages ให้ใช้ fulfillmentText (ถ้ามี)
                if not messages_to_reply and 'fulfillmentText' in response_dict['queryResult']:
                    text_response = response_dict['queryResult']['fulfillmentText']
                    if text_response and text_response not in INVALID_DIALOGFLOW_RESPONSES:
                        text_message = TextMessage(text=text_response, quick_reply=quick_replies if quick_replies else None)
                        messages_to_reply.append(text_message)
                
                # หากมีอย่างน้อยหนึ่งข้อความให้ส่งทั้งหมด
                if messages_to_reply:
                    logger.info("พบคำตอบจาก Dialogflow ส่งคำตอบให้ผู้ใช้")
                    # เพิ่ม quick replies ให้กับข้อความสุดท้าย (ถ้ามี quick replies แต่ยังไม่ได้ใส่)
                    if quick_replies and not messages_to_reply[-1].quick_reply:
                        messages_to_reply[-1].quick_reply = quick_replies
                    
                    send_multiple_messages(event.reply_token, messages_to_reply)
                else:
                    # ขั้นตอนที่ 2: ค้นหาในเอกสาร
                    logger.info("เริ่มขั้นตอนที่ 2: ค้นหาในเอกสาร")
                    reply_text, found_in_docs = search_from_documents(actual_message)

                    if not found_in_docs:
                        # ขั้นตอนที่ 3: ใช้ Hugging Face
                        logger.info("เริ่มขั้นตอนที่ 3: ใช้ Hugging Face")
                        reply_text = ask_huggingface_model(actual_message)
                        logger.info("ได้คำตอบจาก Hugging Face")
                    
                    # ส่งข้อความที่ได้
                    text_message = TextMessage(text=reply_text, quick_reply=quick_replies if quick_replies else None)
                    send_multiple_messages(event.reply_token, [text_message])
                
            except Exception as e:
                logger.error(f"เกิดข้อผิดพลาดในการประมวลผลข้อความ: {str(e)}")
                send_text_message(event.reply_token, "ขออภัย เกิดข้อผิดพลาดในการประมวลผล กรุณาลองใหม่อีกครั้ง")

def process_payload(payload, messages_list):
    """
    ประมวลผล payload และเพิ่มเข้าไปในรายการข้อความ
    โดยเฉพาะ Flex Message ของ LINE
    """
    try:
        logger.info(f"กำลังประมวลผล payload: {json.dumps(payload, indent=2, ensure_ascii=False)[:500]}")
        
        # ตรวจสอบว่ามี LINE Flex Message หรือไม่
        if 'line' in payload and isinstance(payload['line'], dict):
            line_content = payload['line']
            
            if 'type' in line_content and line_content['type'] == 'flex':
                # สร้าง Flex Message
                flex_message = create_flex_message(line_content)
                if flex_message:
                    messages_list.append(flex_message)
                return
        
        # ตรวจสอบกรณีที่ payload เป็น Flex Message โดยตรง
        if isinstance(payload, dict) and 'type' in payload and payload['type'] == 'flex':
            # สร้าง Flex Message
            flex_message = create_flex_message(payload)
            if flex_message:
                messages_list.append(flex_message)
            return
            
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการประมวลผล payload: {str(e)}")

def create_flex_message(flex_content):
    """
    สร้าง Flex Message จาก payload สำหรับ LINE
    """
    try:
        # Debug
        logger.info(f"กำลังสร้าง Flex Message: {json.dumps(flex_content)[:200]}...")
        
        # ตรวจสอบและแก้ไขโครงสร้าง JSON ที่อาจมีปัญหา
        if 'contents' in flex_content:
            # แก้ไขปัญหาที่อาจมีในโครงสร้าง JSON
            flex_contents = flex_content['contents']
            # ทำความสะอาดข้อมูล JSON (ถ้าจำเป็น)
            if isinstance(flex_contents, dict):
                # ตรวจสอบความถูกต้องของ JSON
                json.dumps(flex_contents)  # ทดสอบว่าสามารถแปลงเป็น JSON ได้

            # สร้าง FlexContainer object (จำเป็นสำหรับ LINE API v3)
            flex_container = FlexContainer.from_dict(flex_content['contents'])
            
            # สร้างและส่งคืน FlexMessage object
            return FlexMessage(
                alt_text=flex_content.get('altText', 'Flex Message'),
                contents=flex_container
            )
        
        return None
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการสร้าง Flex Message: {str(e)}")
        return None

def send_multiple_messages(reply_token, messages):
    """
    ส่งหลายข้อความในครั้งเดียว
    """
    try:
        if not messages:
            logger.warning("ไม่มีข้อความที่จะส่ง")
            return
            
        logger.info(f"กำลังส่ง {len(messages)} ข้อความ")
        reply_request = ReplyMessageRequest(
            reply_token=reply_token,
            messages=messages
        )
        line_bot_api.reply_message_with_http_info(reply_request)
        logger.info("ส่งข้อความสำเร็จ")
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการส่งข้อความหลายรายการ: {str(e)}")
        # ส่งข้อความสำรองในกรณีที่เกิดข้อผิดพลาด
        try:
            send_text_message(reply_token, "ขออภัย เกิดข้อผิดพลาดในการส่งข้อความ")
        except:
            logger.error("ไม่สามารถส่งข้อความสำรองได้")

def send_text_message(reply_token, text):
    """
    ส่งข้อความตัวอักษรธรรมดา
    """
    try:
        # ตรวจสอบว่าข้อความไม่เป็นค่าว่าง
        text = text if text else "ขออภัย ไม่พบข้อมูล"
        
        # ตัดความยาวข้อความถ้าเกิน limit ของ LINE (5000 ตัวอักษร)
        if len(text) > 4997:
            text = text[:4997] + "..."
            
        logger.info(f"กำลังส่งข้อความตอบกลับ: {text[:100]}...")
        reply_request = ReplyMessageRequest(
            reply_token=reply_token,
            messages=[TextMessage(text=text)]
        )
        line_bot_api.reply_message_with_http_info(reply_request)
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการส่งข้อความตัวอักษร: {str(e)}")

def detect_intent_texts(project_id, session_id, text, language_code):
    """
    ส่งข้อความไปยัง Dialogflow เพื่อตรวจจับเจตนา (intent)
    """
    try:
        logger.info(f"กำลังติดต่อ Dialogflow: Project={project_id}, Session={session_id}")
        session_client = SessionsClient()
        session = session_client.session_path(project_id, session_id)
        text_input = TextInput(text=text, language_code=language_code)
        query_input = QueryInput(text=text_input)
        response = session_client.detect_intent(request={"session": session, "query_input": query_input})
        return response
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดกับ Dialogflow: {str(e)}")
        # สร้าง response จำลองเพื่อให้โค้ดยังทำงานต่อได้
        class MockResponse:
            class MockQueryResult:
                fulfillment_text = ""
                fulfillment_messages = []
            query_result = MockQueryResult()
            _pb = type('MockPb', (object,), {})()
        return MockResponse()

def extract_text_from_docx(file_path):
    """
    สกัดข้อความจากไฟล์ Word (.docx)
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # ดึงข้อความจากเอกสาร
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
                
        # ดึงข้อความจากตาราง
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)
                        
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์ docx: {file_path}: {str(e)}")
        return ""

def extract_text_from_txt(file_path):
    """
    อ่านข้อความจากไฟล์ .txt
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # ลองอ่านอีกครั้งด้วย encoding อื่น
        try:
            with open(file_path, 'r', encoding='tis-620') as f:  # encoding สำหรับภาษาไทย
                return f.read()
        except:
            logger.error(f"ไม่สามารถอ่านไฟล์ txt ได้: {file_path}")
            return ""
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์ txt: {file_path}: {str(e)}")
        return ""

def extract_data_from_json(file_path):
    """
    อ่านข้อมูลจากไฟล์ .json
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์ json: {file_path}: {str(e)}")
        return {}

def similar(a, b):
    """
    คำนวณความคล้ายคลึงของข้อความสองชุด
    ค่าที่คืนคือตัวเลขระหว่าง 0 ถึง 1 โดยที่ 1 คือเหมือนกันทั้งหมด
    """
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

def get_data_files(directory='data'):
    """
    ค้นหาไฟล์ข้อมูลทั้งหมดในไดเรกทอรีและ subdirectories
    """
    data_files = {
        'docx': [],
        'json': [],
        'txt': [],
    }
    
    # หาไดเรกทอรีหลักของโปรเจค
    base_dir = os.path.abspath(os.path.dirname(__file__))
    data_dir = os.path.join(base_dir, directory)
    
    # ตรวจสอบว่าไดเรกทอรีมีอยู่จริงหรือไม่
    if not os.path.exists(data_dir):
        logger.warning(f"ไม่พบไดเรกทอรี: {data_dir}")
        return data_files
    
    # เดินทางผ่านไดเรกทอรีและเก็บรวบรวมไฟล์
    for root, _, files in os.walk(data_dir):
        for file in files:
            file_path = os.path.join(root, file)
            _, ext = os.path.splitext(file)
            
            if ext.lower() == '.docx':
                data_files['docx'].append(file_path)
            elif ext.lower() == '.json':
                data_files['json'].append(file_path)
            elif ext.lower() == '.txt':
                data_files['txt'].append(file_path)
    
    return data_files

def search_from_documents(question):
    """
    ค้นหาคำตอบจากไฟล์เอกสารหลายประเภท (docx, json, txt)
    """
    try:
        # หาไฟล์ข้อมูลทั้งหมด
        all_files = get_data_files('data')
        logger.info(f"เริ่มค้นหาคำถาม: '{question}' ในเอกสาร")
        logger.info(f"พบไฟล์เอกสารทั้งหมด: docx={len(all_files['docx'])}, json={len(all_files['json'])}, txt={len(all_files['txt'])}")
        
        best_match = None
        best_match_score = 0
        best_match_answer = ""
        found_answer = False
        
        # ค้นหาในไฟล์ JSON
        for json_file in all_files['json']:
            try:
                data = extract_data_from_json(json_file)
                
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict) and "question" in item and "answer" in item:
                            similarity = similar(question, item["question"])
                            if similarity > best_match_score and similarity > 0.7:
                                found_answer = True
                                best_match = item["question"]
                                best_match_score = similarity
                                best_match_answer = item["answer"]
                                logger.info(f"พบคำตอบที่คล้ายกัน ({similarity:.2f}) จากไฟล์ JSON: {json_file}")
                                logger.info(f"คำถามที่คล้าย: {item['question']}")
                                logger.info(f"ค่าความคล้าย: {similarity:.2f}")
            except Exception as e:
                logger.error(f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ JSON {json_file}: {str(e)}")
                continue

        # ค้นหาในไฟล์ DOCX
        if best_match_score < 0.8:
            for docx_file in all_files['docx']:
                try:
                    content = extract_text_from_docx(docx_file)
                    if not content:
                        continue
                    
                    paragraphs = content.split('\n')
                    for i, para in enumerate(paragraphs):
                        if not para.strip():
                            continue
                            
                        similarity = similar(question, para)
                        if similarity > best_match_score and similarity > 0.6:
                            found_answer = True
                            best_match = para
                            best_match_score = similarity
                            best_match_answer = paragraphs[i+1] if i+1 < len(paragraphs) else para
                            logger.info(f"พบคำตอบจากไฟล์ DOCX ({similarity:.2f}): {docx_file}")
                            logger.info(f"ย่อหน้าที่คล้าย: {para[:100]}...")
                            logger.info(f"ค่าความคล้าย: {similarity:.2f}")
                except Exception as e:
                    logger.error(f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ DOCX {docx_file}: {str(e)}")
                    continue

        # ค้นหาในไฟล์ TXT
        if best_match_score < 0.8:
            for txt_file in all_files['txt']:
                try:
                    content = extract_text_from_txt(txt_file)
                    if not content:
                        continue
                    
                    paragraphs = content.split('\n')
                    for i, para in enumerate(paragraphs):
                        if not para.strip():
                            continue
                            
                        similarity = similar(question, para)
                        if similarity > best_match_score and similarity > 0.6:
                            found_answer = True
                            best_match = para
                            best_match_score = similarity
                            best_match_answer = paragraphs[i+1] if i+1 < len(paragraphs) else para
                            logger.info(f"พบคำตอบจากไฟล์ TXT ({similarity:.2f}): {txt_file}")
                            logger.info(f"ย่อหน้าที่คล้าย: {para[:100]}...")
                            logger.info(f"ค่าความคล้าย: {similarity:.2f}")
                except Exception as e:
                    logger.error(f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ TXT {txt_file}: {str(e)}")
                    continue
        
        # ถ้าไม่พบคำตอบจากไฟล์ใด ๆ
        if not found_answer:
            logger.info("ไม่พบคำตอบในเอกสาร")
            return "ขออภัย ฉันไม่พบข้อมูลที่เกี่ยวข้องกับคำถามของคุณในเอกสารของเรา คุณสามารถถามคำถามเกี่ยวกับหัวข้ออื่นได้ค่ะ", False
        
        # ถ้าพบคำตอบ
        logger.info(f"พบคำตอบในเอกสาร (ค่าความคล้าย: {best_match_score:.2f})")
        return best_match_answer.strip(), True
            
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการค้นหาเอกสาร: {str(e)}")
        return "ขออภัย เกิดข้อผิดพลาดในการค้นหาข้อมูล กรุณาลองใหม่อีกครั้ง", False

def ask_huggingface_model(question):
    """
    ถามคำถามกับโมเดล Hugging Face
    """
    try:
        if huggingface_qa_pipeline is None:
            logger.warning("ไม่สามารถใช้งาน Hugging Face model ได้")
            return "ขออภัย ระบบไม่สามารถโหลดโมเดล AI ได้ กรุณาติดต่อผู้ดูแลระบบ"
        
        logger.info(f"กำลังถามโมเดล Hugging Face: '{question}'")
        
        # สร้าง prompt ที่ชัดเจน
        prompt = f"""โปรดตอบคำถามต่อไปนี้อย่างละเอียด:

คำถาม: {question}

คำตอบ: """
        
        # ปรับพารามิเตอร์ให้เหมาะสม
        outputs = huggingface_qa_pipeline(
            prompt,
            max_new_tokens=200,
            num_return_sequences=1,
            no_repeat_ngram_size=3,
            temperature=0.8,
            top_k=50,
            top_p=0.95,
            do_sample=True
        )
        
        # แยกคำตอบและทำความสะอาด
        generated_text = outputs[0]['generated_text']
        try:
            answer = generated_text.split("คำตอบ:", 1)[1].strip()
            # ลบคำซ้ำ
            words = answer.split()
            answer = " ".join(dict.fromkeys(words))
        except:
            answer = generated_text
        
        logger.info(f"ได้รับคำตอบจาก Hugging Face: {answer[:100]}...")
        
        if len(answer) < 10 or any(word in answer.lower() for word in ['error', 'ไม่สามารถ', 'ขออภัย']):
            return "ขออภัย ฉันไม่สามารถตอบคำถามนี้ได้อย่างเหมาะสม กรุณาถามคำถามอื่น หรือถามในรูปแบบอื่น"
            
        return answer

    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดกับ Hugging Face: {str(e)}")
        return "ขออภัย ระบบกำลังมีปัญหาในการประมวลผล กรุณาลองใหม่ภายหลัง"

@app.route("/")
def home():
    """
    หน้าแรกเพื่อตรวจสอบสถานะการทำงานของบอท
    """
    return "LINE Bot with Dialogflow + Document + Hugging Face กำลังทำงาน!"

@app.route("/status")
def status():
    """
    แสดงสถานะของบอทและส่วนประกอบต่าง ๆ ในรูปแบบ JSON
    """
    doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "documents.json")
    cred_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    
    status_report = {
        "dialogflow": {
            "status": "✅ พร้อมใช้งาน" if DIALOGFLOW_PROJECT_ID and os.path.exists(cred_path) else "❌ ยังไม่ได้ตั้งค่า",
            "project_id": DIALOGFLOW_PROJECT_ID,
            "credentials": "พบไฟล์" if os.path.exists(cred_path) else "ไม่พบไฟล์"
        },
        "documents": {
            "status": "✅ พร้อมใช้งาน" if os.path.exists(doc_path) else "❌ ไม่พบไฟล์",
            "path": doc_path,
            "exists": os.path.exists(doc_path)
        },
        "huggingface_model": {
            "status": "✅ พร้อมใช้งาน" if huggingface_qa_pipeline is not None else "❌ ยังไม่ได้โหลด"
        },
        "line_api": {
            "status": "✅ ตั้งค่าแล้ว" if LINE_CHANNEL_ACCESS_TOKEN and LINE_CHANNEL_SECRET else "❌ ยังไม่ได้ตั้งค่า"
        }
    }

    return jsonify({
        "status": "online",
        "components": status_report,
        "timestamp": datetime.now().isoformat()
    })

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)