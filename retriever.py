import os
import json
import logging
import docx
from difflib import SequenceMatcher

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์ docx: {file_path}: {str(e)}")
        return ""

def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='tis-620') as f:
                return f.read()
        except:
            logger.error(f"ไม่สามารถอ่านไฟล์ txt ได้: {file_path}")
            return ""
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์ txt: {file_path}: {str(e)}")
        return ""

def extract_data_from_json(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์ json: {file_path}: {str(e)}")
        return {}

def similar(a, b):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

def get_data_files(directory='data'):
    data_files = {
        'docx': [],
        'json': [],
        'txt': [],
    }
    base_dir = os.path.abspath(os.path.dirname(__file__))
    data_dir = os.path.join(base_dir, directory)
    if not os.path.exists(data_dir):
        logger.warning(f"ไม่พบไดเรกทอรี: {data_dir}")
        return data_files
    for root, _, files in os.walk(data_dir):
        for file in files:
            file_path = os.path.join(root, file)
            _, ext = os.path.splitext(file)
            if ext.lower() == '.docx':
                data_files['docx'].append(file_path)
            elif ext.lower() == '.json':
                data_files['json'].append(file_path)
            elif ext.lower() == '.txt':
                data_files['txt'].append(file_path)
    return data_files

def search_from_documents(question):
    try:
        all_files = get_data_files('data')
        logger.info(f"เริ่มค้นหาคำถาม: '{question}' ในเอกสาร")
        logger.info(f"พบไฟล์เอกสารทั้งหมด: docx={len(all_files['docx'])}, json={len(all_files['json'])}, txt={len(all_files['txt'])}")
        best_match = None
        best_match_score = 0
        best_match_answer = ""
        found_answer = False
        for json_file in all_files['json']:
            try:
                data = extract_data_from_json(json_file)
                if isinstance(data, list):
                    for item in data:
                        if isinstance(item, dict) and "question" in item and "answer" in item:
                            similarity = similar(question, item["question"])
                            if similarity > best_match_score and similarity > 0.7:
                                found_answer = True
                                best_match = item["question"]
                                best_match_score = similarity
                                best_match_answer = item["answer"]
                                logger.info(f"พบคำตอบที่คล้ายกัน ({similarity:.2f}) จากไฟล์ JSON: {json_file}")
                                logger.info(f"คำถามที่คล้าย: {item['question']}")
                                logger.info(f"ค่าความคล้าย: {similarity:.2f}")
            except Exception as e:
                logger.error(f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ JSON {json_file}: {str(e)}")
                continue
        if best_match_score < 0.8:
            for docx_file in all_files['docx']:
                try:
                    content = extract_text_from_docx(docx_file)
                    if not content:
                        continue
                    paragraphs = content.split('\n')
                    for i, para in enumerate(paragraphs):
                        if not para.strip():
                            continue
                        similarity = similar(question, para)
                        if similarity > best_match_score and similarity > 0.6:
                            found_answer = True
                            best_match = para
                            best_match_score = similarity
                            best_match_answer = paragraphs[i+1] if i+1 < len(paragraphs) else para
                            logger.info(f"พบคำตอบจากไฟล์ DOCX ({similarity:.2f}): {docx_file}")
                            logger.info(f"ย่อหน้าที่คล้าย: {para[:100]}...")
                            logger.info(f"ค่าความคล้าย: {similarity:.2f}")
                except Exception as e:
                    logger.error(f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ DOCX {docx_file}: {str(e)}")
                    continue
        if best_match_score < 0.8:
            for txt_file in all_files['txt']:
                try:
                    content = extract_text_from_txt(txt_file)
                    if not content:
                        continue
                    paragraphs = content.split('\n')
                    for i, para in enumerate(paragraphs):
                        if not para.strip():
                            continue
                        similarity = similar(question, para)
                        if similarity > best_match_score and similarity > 0.6:
                            found_answer = True
                            best_match = para
                            best_match_score = similarity
                            best_match_answer = paragraphs[i+1] if i+1 < len(paragraphs) else para
                            logger.info(f"พบคำตอบจากไฟล์ TXT ({similarity:.2f}): {txt_file}")
                            logger.info(f"ย่อหน้าที่คล้าย: {para[:100]}...")
                            logger.info(f"ค่าความคล้าย: {similarity:.2f}")
                except Exception as e:
                    logger.error(f"เกิดข้อผิดพลาดในการประมวลผลไฟล์ TXT {txt_file}: {str(e)}")
                    continue
        if not found_answer:
            logger.info("ไม่พบคำตอบในเอกสาร")
            return "ขออภัย ฉันไม่พบข้อมูลที่เกี่ยวข้องกับคำถามของคุณในเอกสารของเรา คุณสามารถถามคำถามเกี่ยวกับหัวข้ออื่นได้ค่ะ", False
        logger.info(f"พบคำตอบในเอกสาร (ค่าความคล้าย: {best_match_score:.2f})")
        return best_match_answer.strip(), True
    except Exception as e:
        logger.error(f"เกิดข้อผิดพลาดในการค้นหาเอกสาร: {str(e)}")
        return "ขออภัย เกิดข้อผิดพลาดในการค้นหาข้อมูล กรุณาลองใหม่อีกครั้ง", False
